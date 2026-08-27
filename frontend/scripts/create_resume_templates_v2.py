#!/usr/bin/env python3
"""Build RoleVault's category-specific ATS DOCX skeletons and selector previews."""

from __future__ import annotations

import sys
from dataclasses import dataclass
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SERVER_TEMPLATE_DIR = ROOT / "src" / "server" / "templates"
PUBLIC_TEMPLATE_DIR = ROOT / "public" / "templates"
PREVIEW_DIR = PUBLIC_TEMPLATE_DIR / "previews"


@dataclass(frozen=True)
class TemplateProfile:
    key: str
    accent: str
    title: str
    summary_heading: str
    skills_heading: str
    experience_heading: str
    projects_heading: str | None
    credentials_heading: str
    education_heading: str
    header_alignment: str = "left"
    page_format: str = "letter"
    section_order: tuple[str, ...] = ("summary", "skills", "experience", "projects", "credentials", "education")


PROFILES = (
    TemplateProfile(
        key="finance_accounting",
        accent="17365D",
        title="Finance & Accounting",
        summary_heading="FINANCE PROFILE",
        skills_heading="FINANCE CAPABILITIES & SYSTEMS",
        experience_heading="PROFESSIONAL EXPERIENCE",
        projects_heading="SELECTED TRANSACTIONS & PROJECTS",
        credentials_heading="CERTIFICATIONS & CREDENTIALS",
        education_heading="EDUCATION",
    ),
    TemplateProfile(
        key="sales_marketing",
        accent="7A2E4D",
        title="Sales & Marketing",
        summary_heading="COMMERCIAL PROFILE",
        skills_heading="SALES & MARKETING CAPABILITIES",
        experience_heading="COMMERCIAL EXPERIENCE",
        projects_heading="CAMPAIGNS & SELECTED WORK",
        credentials_heading="CERTIFICATIONS",
        education_heading="EDUCATION",
    ),
    TemplateProfile(
        key="legal",
        accent="315B7D",
        title="Legal",
        summary_heading="LEGAL PROFILE",
        skills_heading="LEGAL CAPABILITIES",
        experience_heading="LEGAL & PROFESSIONAL EXPERIENCE",
        projects_heading=None,
        credentials_heading="ADMISSION, PLT & CREDENTIALS",
        education_heading="EDUCATION",
        page_format="a4",
        section_order=("summary", "education", "credentials", "experience", "skills"),
    ),
    TemplateProfile(
        key="human_resources_admin_operations",
        accent="245B4A",
        title="People & Operations",
        summary_heading="PEOPLE & OPERATIONS PROFILE",
        skills_heading="CORE EXPERTISE",
        experience_heading="PROFESSIONAL EXPERIENCE",
        projects_heading="SELECTED PEOPLE & OPERATIONS INITIATIVES",
        credentials_heading="PROFESSIONAL CREDENTIALS",
        education_heading="EDUCATION",
    ),
    TemplateProfile(
        key="hospitality_retail_customer_service",
        accent="74452D",
        title="Service & Hospitality",
        summary_heading="SERVICE PROFILE",
        skills_heading="SERVICE CAPABILITIES",
        experience_heading="SERVICE EXPERIENCE",
        projects_heading=None,
        credentials_heading="LICENCES & TRAINING",
        education_heading="EDUCATION",
        header_alignment="center",
    ),
    TemplateProfile(
        key="general_professional_other",
        accent="244A73",
        title="General Professional",
        summary_heading="PROFESSIONAL SUMMARY",
        skills_heading="CORE SKILLS",
        experience_heading="PROFESSIONAL EXPERIENCE",
        projects_heading=None,
        credentials_heading="CERTIFICATIONS & LICENCES",
        education_heading="EDUCATION",
    ),
)


def set_font(style, size: float, *, bold: bool = False, color: str = "20242A") -> None:
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), "Arial")


def set_keep(paragraph, *, next_paragraph: bool = False, together: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if next_paragraph:
        p_pr.append(OxmlElement("w:keepNext"))
    if together:
        p_pr.append(OxmlElement("w:keepLines"))


def add_bottom_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(document: Document, profile: TemplateProfile) -> None:
    section = document.sections[0]
    if profile.page_format == "a4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = document.styles
    normal = styles["Normal"]
    set_font(normal, 9.8)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.04

    name = styles.add_style("Resume Name", WD_STYLE_TYPE.PARAGRAPH)
    set_font(name, 23, bold=True, color=profile.accent)
    name.paragraph_format.space_after = Pt(1)
    name.paragraph_format.keep_with_next = True

    role = styles.add_style("Resume Role", WD_STYLE_TYPE.PARAGRAPH)
    set_font(role, 11.2, bold=True, color="323942")
    role.paragraph_format.space_after = Pt(2)
    role.paragraph_format.keep_with_next = True

    contact = styles.add_style("Resume Contact", WD_STYLE_TYPE.PARAGRAPH)
    set_font(contact, 9.1, color="59616B")
    contact.paragraph_format.space_after = Pt(7)

    section_heading = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    set_font(section_heading, 10.2, bold=True, color=profile.accent)
    section_heading.paragraph_format.space_before = Pt(7)
    section_heading.paragraph_format.space_after = Pt(4)
    section_heading.paragraph_format.keep_with_next = True

    item_heading = styles.add_style("Resume Item", WD_STYLE_TYPE.PARAGRAPH)
    set_font(item_heading, 9.8, bold=True, color="20242A")
    item_heading.paragraph_format.space_before = Pt(4)
    item_heading.paragraph_format.space_after = Pt(0)
    item_heading.paragraph_format.keep_with_next = True

    item_meta = styles.add_style("Resume Item Meta", WD_STYLE_TYPE.PARAGRAPH)
    set_font(item_meta, 9.2, color="535B65")
    item_meta.paragraph_format.space_after = Pt(1.5)
    item_meta.paragraph_format.keep_with_next = True

    bullet = styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    set_font(bullet, 9.5)
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.space_after = Pt(1.5)
    bullet.paragraph_format.line_spacing = 1.03

    control = styles.add_style("Template Control", WD_STYLE_TYPE.PARAGRAPH)
    set_font(control, 1, color="FFFFFF")
    control.paragraph_format.space_before = Pt(0)
    control.paragraph_format.space_after = Pt(0)
    control.paragraph_format.line_spacing = 0.01

    document.core_properties.title = f"RoleVault {profile.title} Resume Template v2"
    document.core_properties.subject = "ATS-friendly one-column resume template"
    document.core_properties.author = "RoleVault"
    document.core_properties.comments = (
        f"Compact Reference Guide preset with RoleVault resume overrides: {profile.page_format.upper()}, "
        "Arial, 0.65-inch side margins, 0.55-inch vertical margins, one column, no tables."
    )


def add_control(document: Document, tag: str) -> None:
    paragraph = document.add_paragraph(style="Template Control")
    paragraph.add_run(tag)


def add_heading(document: Document, text: str, accent: str) -> None:
    paragraph = document.add_paragraph(text, style="Resume Section")
    add_bottom_border(paragraph, accent)
    set_keep(paragraph, next_paragraph=True)


def add_simple_section(document: Document, heading: str, body_tag: str, accent: str) -> None:
    add_heading(document, heading, accent)
    paragraph = document.add_paragraph(body_tag)
    set_keep(paragraph, together=True)


def add_conditional_simple_section(
    document: Document, condition: str, heading: str, body_tag: str, accent: str
) -> None:
    add_control(document, f"{{#{condition}}}")
    add_simple_section(document, heading, body_tag, accent)
    add_control(document, f"{{/{condition}}}")


def add_experience_section(document: Document, heading: str, accent: str) -> None:
    add_control(document, "{#hasExperience}")
    add_heading(document, heading, accent)
    add_control(document, "{/hasExperience}")
    add_control(document, "{#experience}")
    paragraph = document.add_paragraph(style="Resume Item")
    paragraph.add_run("{title}")
    set_keep(paragraph, next_paragraph=True)
    paragraph = document.add_paragraph("{experienceMetaLine}", style="Resume Item Meta")
    set_keep(paragraph, next_paragraph=True)
    add_control(document, "{#bullets}")
    paragraph = document.add_paragraph(style="Resume Bullet")
    paragraph.add_run("•  {.}")
    set_keep(paragraph, together=True)
    add_control(document, "{/bullets}")
    add_control(document, "{/experience}")


def add_projects_section(document: Document, heading: str, accent: str) -> None:
    add_control(document, "{#hasProjects}")
    add_heading(document, heading, accent)
    add_control(document, "{/hasProjects}")
    add_control(document, "{#projects}")
    paragraph = document.add_paragraph(style="Resume Item")
    paragraph.add_run("{projectHeadingLine}")
    set_keep(paragraph, next_paragraph=True)
    add_control(document, "{#bullets}")
    paragraph = document.add_paragraph(style="Resume Bullet")
    paragraph.add_run("•  {.}")
    set_keep(paragraph, together=True)
    add_control(document, "{/bullets}")
    add_control(document, "{/projects}")


def add_credentials_section(document: Document, heading: str, accent: str) -> None:
    add_control(document, "{#hasCredentials}")
    add_heading(document, heading, accent)
    add_control(document, "{/hasCredentials}")
    add_control(document, "{#credentials}")
    paragraph = document.add_paragraph(style="Resume Item")
    paragraph.add_run("{credentialLine}")
    set_keep(paragraph, together=True)
    add_control(document, "{/credentials}")


def add_education_section(document: Document, heading: str, accent: str) -> None:
    add_control(document, "{#hasEducation}")
    add_heading(document, heading, accent)
    add_control(document, "{/hasEducation}")
    add_control(document, "{#education}")
    paragraph = document.add_paragraph(style="Resume Item")
    paragraph.add_run("{educationLine}")
    set_keep(paragraph, next_paragraph=True)
    add_control(document, "{#hasDetails}")
    paragraph = document.add_paragraph(style="Resume Bullet")
    paragraph.add_run("•  {educationDetailsLine}")
    set_keep(paragraph, together=True)
    add_control(document, "{/hasDetails}")
    add_control(document, "{/education}")


def build_template(profile: TemplateProfile) -> Path:
    document = Document()
    configure_document(document, profile)

    align = (
        WD_ALIGN_PARAGRAPH.CENTER
        if profile.header_alignment == "center"
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    for value, style in (
        ("{fullName}", "Resume Name"),
        ("{professionalTitle}", "Resume Role"),
        ("{contactLine}", "Resume Contact"),
    ):
        paragraph = document.add_paragraph(value, style=style)
        paragraph.alignment = align

    section_builders = {
        "summary": lambda: add_simple_section(
            document, profile.summary_heading, "{professionalSummary}", profile.accent
        ),
        "skills": lambda: add_conditional_simple_section(
            document, "hasSkills", profile.skills_heading, "{skillsLine}", profile.accent
        ),
        "experience": lambda: add_experience_section(
            document, profile.experience_heading, profile.accent
        ),
        "projects": lambda: (
            add_projects_section(document, profile.projects_heading, profile.accent)
            if profile.projects_heading else None
        ),
        "credentials": lambda: add_credentials_section(
            document, profile.credentials_heading, profile.accent
        ),
        "education": lambda: add_education_section(
            document, profile.education_heading, profile.accent
        ),
    }
    for section_name in profile.section_order:
        section_builders[section_name]()

    output = SERVER_TEMPLATE_DIR / f"{profile.key}_v2.docx"
    document.save(output)
    return output


def build_preview(profile: TemplateProfile) -> Path:
    heading_by_section = {
        "summary": profile.summary_heading,
        "skills": profile.skills_heading,
        "experience": profile.experience_heading,
        "projects": profile.projects_heading,
        "credentials": profile.credentials_heading,
        "education": profile.education_heading,
    }
    headings = [heading_by_section[section_name] for section_name in profile.section_order]
    headings = [heading for heading in headings if heading]
    bars: list[str] = []
    y = 105
    for index, _heading in enumerate(headings[:5]):
        bars.append(
            f'<rect x="42" y="{y}" width="{118 + (index % 2) * 24}" height="7" rx="2" fill="#{profile.accent}"/>'
        )
        bars.append(
            f'<rect x="42" y="{y + 15}" width="250" height="4" rx="2" fill="#d9dde3"/>'
        )
        bars.append(
            f'<rect x="42" y="{y + 25}" width="{225 - (index % 3) * 18}" height="4" rx="2" fill="#d9dde3"/>'
        )
        if index >= 2:
            bars.append(
                f'<circle cx="47" cy="{y + 38}" r="2" fill="#{profile.accent}"/>'
            )
            bars.append(
                f'<rect x="54" y="{y + 36}" width="226" height="4" rx="2" fill="#e4e7eb"/>'
            )
        y += 64

    title_x = 170 if profile.header_alignment == "center" else 42
    anchor = "middle" if profile.header_alignment == "center" else "start"
    escaped_title = escape(profile.title)
    escaped_description_title = escape(profile.title.lower())
    svg = f'''<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 340 440" role="img" aria-labelledby="title description">
  <title id="title">{escaped_title} resume preview</title>
  <desc id="description">A one-column ATS-friendly {escaped_description_title} resume with discipline-specific sections.</desc>
  <rect width="340" height="440" rx="8" fill="#f5f4f0"/>
  <rect x="18" y="18" width="304" height="404" rx="4" fill="#ffffff" stroke="#c9c6bd"/>
  <text x="{title_x}" y="52" text-anchor="{anchor}" font-family="Arial, sans-serif" font-size="17" font-weight="700" fill="#{profile.accent}">ALEX MORGAN</text>
  <rect x="42" y="65" width="118" height="6" rx="2" fill="#505761"/>
  <rect x="42" y="82" width="238" height="4" rx="2" fill="#a9afb7"/>
  {''.join(bars)}
</svg>
'''
    output = PREVIEW_DIR / f"{profile.key}_v2.svg"
    output.write_text(svg, encoding="utf-8")
    return output


def main() -> None:
    requested = set(sys.argv[1:])
    known = {profile.key for profile in PROFILES}
    unknown = requested - known
    if unknown:
        options = ", ".join(sorted(known))
        raise SystemExit(
            f"Unknown template profile(s): {', '.join(sorted(unknown))}. Choose from: {options}"
        )

    SERVER_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)
    for profile in PROFILES:
        if requested and profile.key not in requested:
            continue
        print(build_template(profile))
        print(build_preview(profile))


if __name__ == "__main__":
    main()
