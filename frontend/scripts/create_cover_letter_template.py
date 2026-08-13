from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "public" / "templates" / "cover_letter_v1.docx"
BLUE = RGBColor(0x0D, 0x38, 0x80)
INK = RGBColor(0x18, 0x1D, 0x26)
MUTED = RGBColor(0x5F, 0x66, 0x70)


def set_cell_border(cell, color="D9DDE5", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_text(paragraph, text, size=10.5, bold=False, color=INK):
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def main():
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.08

    header = document.add_table(rows=1, cols=1)
    header.autofit = True
    cell = header.cell(0, 0)
    set_cell_border(cell, color="0D3880", size="14")
    cell.margin_top = 0
    cell.margin_bottom = 0
    name = cell.paragraphs[0]
    name.paragraph_format.space_after = Pt(2)
    add_text(name, "{candidateName}", size=18, bold=True, color=BLUE)
    contact = cell.add_paragraph()
    contact.paragraph_format.space_after = Pt(8)
    add_text(contact, "{contactLine}", size=9.25, color=MUTED)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(5)

    date = document.add_paragraph()
    date.paragraph_format.space_after = Pt(10)
    add_text(date, "{date}", size=10, color=MUTED)

    recipient = document.add_paragraph()
    recipient.paragraph_format.space_after = Pt(14)
    add_text(recipient, "{recipientBlock}", size=10)

    salutation = document.add_paragraph()
    salutation.paragraph_format.space_after = Pt(9)
    add_text(salutation, "{salutation}", bold=True)

    opening = document.add_paragraph()
    add_text(opening, "{openingParagraph}")

    evidence_start = document.add_paragraph()
    add_text(evidence_start, "{#bodyParagraphs}")
    evidence = document.add_paragraph()
    add_text(evidence, "{text}")
    evidence_end = document.add_paragraph()
    add_text(evidence_end, "{/bodyParagraphs}")

    closing = document.add_paragraph()
    add_text(closing, "{closingParagraph}")

    signoff = document.add_paragraph()
    signoff.paragraph_format.space_before = Pt(5)
    signoff.paragraph_format.space_after = Pt(13)
    add_text(signoff, "{signOff}")

    signature = document.add_paragraph()
    signature.paragraph_format.space_after = Pt(0)
    add_text(signature, "{candidateName}", bold=True, color=BLUE)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
